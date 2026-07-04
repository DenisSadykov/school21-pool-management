from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUTPUT_DIR = Path("/Users/denissadykov/school21-pool-management/docs/artifacts")
OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
OUTPUT_PATH = OUTPUT_DIR / "school21_exam_volunteer_brief.docx"


BLACK = RGBColor(0, 0, 0)
GRAY = RGBColor(85, 85, 85)
BLUE = RGBColor(46, 116, 181)


def set_run_font(run, name="Arial", size=11, bold=False, color=BLACK):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    run.bold = bold
    run.font.color.rgb = color


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def add_bullet(doc, text, style="List Bullet"):
    p = doc.add_paragraph(style=style)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.15
    run = p.add_run(text)
    set_run_font(run, size=11)
    return p


def add_numbered(doc, text):
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = 1.15
    run = p.add_run(text)
    set_run_font(run, size=11)
    return p


def add_body(doc, text, bold_prefix=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.15
    if bold_prefix:
        prefix = p.add_run(bold_prefix)
        set_run_font(prefix, size=11, bold=True)
        rest = p.add_run(text)
        set_run_font(rest, size=11)
    else:
        run = p.add_run(text)
        set_run_font(run, size=11)
    return p


def add_section_heading(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(11)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    set_run_font(run, size=13, bold=True, color=BLUE)
    return p


def configure_styles(doc):
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    normal = doc.styles["Normal"]
    normal.font.name = "Arial"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    normal.font.size = Pt(11)

    for style_name in ("List Bullet", "List Number"):
        style = doc.styles[style_name]
        style.font.name = "Arial"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
        style.font.size = Pt(11)


def build_doc():
    doc = Document()
    configure_styles(doc)

    title = doc.add_paragraph()
    title.paragraph_format.space_after = Pt(3)
    title_run = title.add_run("Бриф волонтёра: проведение экзамена в Школе 21")
    set_run_font(title_run, size=22, bold=True)

    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(8)
    subtitle_run = subtitle.add_run(
        "Готовая памятка на день экзамена для встречающих и кластерных волонтёров"
    )
    set_run_font(subtitle_run, size=11, color=GRAY)

    info_table = doc.add_table(rows=1, cols=3)
    info_table.autofit = False
    widths = [Inches(2.1), Inches(2.1), Inches(2.3)]
    labels = [
        ("Роль", "Волонтёр экзамена"),
        ("Задача", "Порядок, поддержка, навигация"),
        ("Главный принцип", "Спокойно, чётко, по правилам"),
    ]
    for idx, cell in enumerate(info_table.rows[0].cells):
        cell.width = widths[idx]
        set_cell_shading(cell, "F2F4F7")
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        label_run = p.add_run(labels[idx][0] + "\n")
        set_run_font(label_run, size=9, bold=True, color=GRAY)
        value_run = p.add_run(labels[idx][1])
        set_run_font(value_run, size=10, bold=True)

    add_section_heading(doc, "1. Перед началом экзамена")
    add_numbered(doc, "Уточни своё распределение: кластер, зона ответственности и имя лида.")
    add_numbered(doc, "Проверь явку своей команды и сразу сообщи лиду, если кого-то нет или кто-то опаздывает.")
    add_numbered(doc, "Освежи в памяти общий сценарий экзамена и ключевой тайминг.")

    add_section_heading(doc, "2. Подготовка кластера до запуска")
    bullets = [
        "Собрать забытые вещи и отнести их в штаб.",
        "Проверить пространство под клавиатурами и рядом с рабочими местами на наличие шпаргалок.",
        "Убедиться, что все компьютеры залиты и готовы к работе; если есть проблема, сразу подсветить support.",
        "Проверить наличие флешек в APMax и при необходимости отнести их в штаб.",
        "Стереть все надписи с досок.",
        "Осмотреть кабинки и туалеты на наличие шпаргалок, включая труднодоступные места.",
        "Ровно расставить стулья у каждого компьютера.",
        "Разложить на каждое место конверт для личных вещей; на последний экзамен подготовить также наклейки.",
        "Подготовить отдельный стол с листами и ручками для пиров.",
    ]
    for item in bullets:
        add_bullet(doc, item)

    add_section_heading(doc, "3. Запуск, рассадка и старт экзамена")
    add_body(
        doc,
        "Находимся на своих позициях согласно распределению и спокойно навигируем участников.",
        bold_prefix="Общий принцип: ",
    )
    start_bullets = [
        "На турникетах: если браслета нет, выдаём предупреждение и надеваем зелёный браслет. Если браслет рвётся или почти порван, меняем на браслет того же цвета.",
        "Рассадка: участник сначала занимает место, которое называет волонтёр, а затем идёт по делам: гардероб, туалет, вода, сдача личных вещей в конверт A5 или A4 по требованию.",
        "Рассаживаем по часовой стрелке. Кластер заполняется последовательно; после заполнения Galaxy кластер Oxygen перекрывается и запускается Void.",
        "На первый экзамен участник приносит с собой развёрнутый белый лист с тремя шпаргалками: адрес сайта экзамена, логин и пароль от платформы, подсказку по генерации ssh-ключа, а также ручку и бутылку воды с закручивающейся крышкой объёмом до 1,5 л.",
        "Логиниться участники начинают только после официального объявления старта. Если кто-то вошёл раньше, просим разлогиниться.",
    ]
    for item in start_bullets:
        add_bullet(doc, item)

    add_section_heading(doc, "4. Тайминг перед стартом экзамена")
    timing_bullets = [
        "За 10 минут до экзамена все участники должны занять свои места.",
        "Волонтёры проверяют шпаргалки. Если есть лишнее, шпаргалка забирается. Если участник успевает, он может под контролем волонтёра переписать допустимую шпаргалку по правилам. Также проверяем качество запечатанных конвертов и при необходимости меняем их.",
        "За 5 минут до начала один из волонтёров громко и чётко озвучивает правила на весь кластер. Если кластер большой, правила повторяются дважды на две половины.",
        "За 1 минуту до старта сообщаем, что осталась одна минута, и просим кластер замолчать.",
        "В момент старта один волонтёр в кластере вслух объявляет: «Экзамен начался!»",
        "Сразу после старта выдаём ручки и листы A4 тем, кому это требуется.",
    ]
    for item in timing_bullets:
        add_bullet(doc, item)

    add_section_heading(doc, "5. Нештатные ситуации после старта")
    incident_bullets = [
        "Опоздавший участник: запускаем ровно до 12:50. Встречающие волонтёры остаются на первом этаже до 13:30. Если время прошло, спокойно объясняем, что допустить на экзамен уже нельзя, поддерживаем участника и помогаем ему сохранить спокойствие.",
        "Спор по регистрации: если support подтвердил, что регистрации нет, мы не можем самостоятельно записать участника или изменить данные. Спокойно объясняем это и просим внимательнее проверить регистрацию в следующий раз.",
        "Участник уверен, что решил задание правильно, но ответ не принят: объясняем, что это типичная ситуация, просим проверить граничные случаи и при необходимости советуем обратиться к участнику, уже успешно сдавшему задание.",
        "Любая спорная ситуация или отказ следовать указанию волонтёра: сразу зовём TL волонтёров, кратко объясняем ситуацию и передаём решение ему.",
        "Ни при каких обстоятельствах не грубим, не спорим с улыбкой наперекор участнику и не принимаем самовольных решений, которые нарушают правила экзамена.",
    ]
    for item in incident_bullets:
        add_bullet(doc, item)

    add_section_heading(doc, "6. Правила поведения во время экзамена")
    add_body(
        doc,
        "С момента старта участникам запрещено переговариваться, переглядываться, подавать друг другу сигналы и знаки, перемещаться по кластеру без разрешения, а также звать волонтёра голосом или жестами.",
    )
    add_body(
        doc,
        "При первом нарушении на первом экзамене выносится предупреждение. Следующее нарушение ведёт к исключению. Таблица-напоминание о нарушениях выводится на заставках iMac.",
    )

    add_section_heading(doc, "7. Туалетный вопрос и сопровождение")
    toilet_bullets = [
        "Во время похода в туалет локскрин делать не нужно.",
        "Волонтёр отводит участника в туалет и остаётся ждать рядом, пока пир не выйдет.",
        "Если участник долго не выходит, корректно уточняем, всё ли у него в порядке и не нужна ли помощь.",
        "Нельзя допускать, чтобы в туалете одновременно находились два и более участника. При необходимости используем уборную для инвалидов.",
    ]
    for item in toilet_bullets:
        add_bullet(doc, item)

    add_section_heading(doc, "8. Кофе-брейк")
    add_bullet(doc, "До 30 минут на одного человека. Лид кластера распределяет поток и следит, чтобы все успели перекусить в свой слот.")

    add_section_heading(doc, "9. Окончание экзамена")
    end_bullets = [
        "Собрать все конверты и ручки, принести их в 307.",
        "Проверить рабочие места и компьютеры после завершения экзамена.",
        "Собраться на дебриф сразу после закрытия процесса.",
    ]
    for item in end_bullets:
        add_bullet(doc, item)

    add_section_heading(doc, "10. Типичные ошибки и технические сбои")
    error_bullets = [
        "Перед экзаменом ещё раз прочитать файл с типичными ошибками и задать вопросы лидy кластера или на общем брифе.",
        "Если во время экзамена произошёл массовый сбой, например у всех зависают автотесты или не грузится сайт, не паникуем.",
        "О проблеме сообщаем сотрудникам support через соответствующие треды и далее действуем по их указанию.",
        "В support обращаемся только после того, как попробовали решить проблему всеми доступными базовыми способами на месте.",
    ]
    for item in error_bullets:
        add_bullet(doc, item)

    add_section_heading(doc, "11. Если участнику стало плохо")
    health_bullets = [
        "Незамедлительно сообщить тимлиду волонтёров.",
        "Вывести участника к АДМ в 307 и позвать сотрудника АДМ.",
        "Если ситуация серьёзная, сразу вызываем скорую помощь.",
        "Никаких своих таблеток и медицинских советов не даём.",
        "Если у участника есть собственные медикаменты, сопровождаем его к личным вещам.",
    ]
    for item in health_bullets:
        add_bullet(doc, item)

    add_section_heading(doc, "12. Коммуникация и экипировка")
    add_bullet(doc, "Вся рабочая коммуникация ведётся в рокет-чате и строго в соответствующих тредах.")
    add_bullet(doc, "Перед стартом выдать волонтёрам футболки и рации.")

    add_section_heading(doc, "Короткое напоминание волонтёру")
    reminder_bullets = [
        "Ты держишь порядок и спокойствие, а не вступаешь в спор.",
        "Если сомневаешься, зови TL или support, а не принимай спорное решение в одиночку.",
        "Каждое действие должно быть понятным участнику и соответствовать правилам экзамена.",
    ]
    for item in reminder_bullets:
        add_bullet(doc, item)

    doc.save(OUTPUT_PATH)


if __name__ == "__main__":
    build_doc()
    print(OUTPUT_PATH)
